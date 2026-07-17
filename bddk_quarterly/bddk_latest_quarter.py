#!/usr/bin/env python3
"""
BDDK Quarterly Loan-Breakdown Extractor — LATEST quarter, 25 banks
=================================================================

Recurring companion to the historical backfill script. Each run:

  1. Auto-detects the most recent quarter for which any of the 25 tracked
     banks has published a consolidated ("Konsolide") report on BDDK's
     "Bağımsız Denetim Raporları" portal (https://www.bddk.gov.tr/BdrUyg).
  2. Downloads that quarter's consolidated report for each of the 25 banks.
  3. Extracts BOTH loan-breakdown footnote tables from each report.
  4. Writes a single-quarter workbook with ONE SHEET PER TABLE, in the
     Template.xlsx layout:
        output/BDDK_25Banks_<QUARTER>.xlsx    (e.g. ..._2026Q2.xlsx)

     sheet "Tüketici Kredileri"        sheet "Taksitli Ticari Krediler"
       Tüketici Kredileri-TP             Taksitli Ticari Krediler-TP
         Konut / Taşıt / İhtiyaç / Diğer   İşyeri / Taşıt / İhtiyaç / Diğer
       Tüketici Kredileri-YP             Taksitli Ticari Krediler-YP
         ... (same four)                   ... (same four)

You never edit this file between quarters — just re-run it. When 2026-Q2 is
published, `python3 bddk_latest_quarter.py` picks it up automatically; same
for 2026-Q3, etc. To pin a specific quarter without editing code:
        python3 bddk_latest_quarter.py --quarter 2026Q2

Only the **Cari Dönem** (current period) table is ever used; each report also
contains an Önceki Dönem (previous period) copy, explicitly excluded.

Extraction per bank/quarter/table, in order:
  * MANUAL_OVERRIDES  — a few verified historical cells.
  * text / .docx parse — the normal path for almost every bank.
  * OCR fallback      — when a bank renders its table as an image / vector
    curves / letter-spaced text (e.g. Fibabanka), the page is rasterized and
    read with RapidOCR (pip-only, no system deps). OCR numbers are accepted
    ONLY if they pass the sum-check (sub-items == total); otherwise the cell
    is left 0 and flagged, with the page image saved to _ocr_pages/.
  * else 0 (no such note, or file missing on BDDK's server).

Rules: no consolidated report for the quarter -> 0; blank / "-" field -> 0.

Dependencies (one-time):
    pip install -r requirements.txt
No Tesseract / Homebrew required — OCR is pure-pip via RapidOCR.
"""

import argparse
import io
import json
import os
import re
import sys
import time
import urllib.parse
import zipfile
from datetime import datetime

import openpyxl
import pdfplumber
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

try:
    import docx  # python-docx; a few banks (e.g. TEB) file .docx reports
except ImportError:
    docx = None

try:
    # bddk.gov.tr serves an incomplete certificate chain; truststore routes
    # verification through the OS trust store (certifi's static bundle fails).
    import truststore
    truststore.inject_into_ssl()
except ImportError:
    pass

BASE = "https://www.bddk.gov.tr/BdrUyg"
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
    )
}

# --------------------------------------------------------------------------
# Table registry (kept in sync with ../bddk_consolidated_loans.py)
# --------------------------------------------------------------------------
FIELD_KEYS = (["TP_Toplam"] + [f"TP_S{i}" for i in range(1, 5)] +
              ["YP_Toplam"] + [f"YP_S{i}" for i in range(1, 5)])
ZERO_RESULT = {k: 0 for k in FIELD_KEYS}

_DASH = r"[-–—]"

# Labels are matched against a SPACE-STRIPPED key (see label_key), so word
# spacing is a non-issue: "Konut Kredisi" (normal), "KonutKredisi" (tight
# kerning defeats pdfplumber's space detection) and "Kre dileri" (RapidOCR
# splitting a word across two boxes) all normalize to the same key. Numbers are
# still read from the original line.
# Also tolerated: case; ş/s, ı/i, ç/c, ğ/g (OCR drops Turkish diacritics);
# singular vs plural ("Konut Kredisi" vs Garanti's "Taşıt Kredileri");
# "İşyeri" vs Garanti's two-word "İş Yeri"; HSBC's "Otomobil Kredisi".
_WS_RE = re.compile(r"\s+")
_NL = r"(?![^\W\d_])"  # not followed by a letter (label must end here)
_KREDI = r"Kredi(?:si|leri)"


# Footnote markers are sometimes rendered inside the label word itself, e.g.
# Emlak Katılım's "Konut Kredis4i". Only digits sandwiched between letters are
# dropped, so the row's actual figures (bounded by spaces/punctuation) survive.
_FOOTNOTE_DIGIT_RE = re.compile(r"(?<=[^\W\d_])\d+(?=[^\W\d_])")


def label_key(line):
    """Space-stripped form of a line, used only for label matching."""
    return _FOOTNOTE_DIGIT_RE.sub("", _WS_RE.sub("", line))


# A row whose label is long can wrap, leaving its figures alone on the next
# line (e.g. Vakıf Katılım's "Tüketici Kredileri-TP" / "54.259 9.954.056
# 10.008.315"). Only a line that is *nothing but* figures may be borrowed, so
# we can never steal the next labelled row's numbers.
_NUMBERS_ONLY_RE = re.compile(r"^[\d.,()\s\-–—]+$")


def row_value(lines, idx):
    """The row's Toplam figure: last number on the line, or on the following
    figures-only continuation line if the label line carries none."""
    v = line_last_number(lines[idx])
    if v is None and idx + 1 < len(lines):
        nxt = lines[idx + 1].strip()
        if nxt and _NUMBERS_ONLY_RE.match(nxt):
            v = line_last_number(nxt)
    return v


KONUT_RE = re.compile(r"^Konut" + _KREDI + _NL, re.IGNORECASE)
ISYERI_RE = re.compile(r"^[İIi][şs][Yy]eri" + _KREDI + _NL, re.IGNORECASE)
TASIT_RE = re.compile(r"^(?:Ta[şs][ıi]t|Otomobil)" + _KREDI + _NL, re.IGNORECASE)
IHTIYAC_RE = re.compile(r"^[İIi]htiya[çc]" + _KREDI + _NL, re.IGNORECASE)
DIGER_RE = re.compile(r"^Di[ğg]er" + _NL, re.IGNORECASE)

TABLE_SPECS = {
    "tuketici": {
        "sheet": "Tüketici Kredileri",
        "hdr_tp": re.compile(r"^T\w*keticiKredileri" + _DASH + r"TP" + _NL, re.IGNORECASE),
        "hdr_yp": re.compile(r"^T\w*keticiKredileri" + _DASH + r"YP" + _NL, re.IGNORECASE),
        "label_tp": "Tüketici Kredileri-TP",
        "label_yp": "Tüketici Kredileri-YP",
        "subs": [("Konut Kredisi", KONUT_RE), ("Taşıt Kredisi", TASIT_RE),
                 ("İhtiyaç Kredisi", IHTIYAC_RE), ("Diğer", DIGER_RE)],
        "section": (re.compile(r"t[üu]ketici\s+kredileri", re.IGNORECASE),
                    re.compile(r"bireysel\s+kredi\s+kart", re.IGNORECASE)),
    },
    "taksitli": {
        "sheet": "Taksitli Ticari Krediler",
        # Banks differ: "Krediler-TP" (most), "Kredileri-TP" (ICBC),
        # "Krediler - TP" (HSBC, spaced dash).
        "hdr_tp": re.compile(r"^TaksitliTicariKredi(?:ler|leri)" + _DASH + r"TP" + _NL, re.IGNORECASE),
        "hdr_yp": re.compile(r"^TaksitliTicariKredi(?:ler|leri)" + _DASH + r"YP" + _NL, re.IGNORECASE),
        "label_tp": "Taksitli Ticari Krediler-TP",
        "label_yp": "Taksitli Ticari Krediler-YP",
        "subs": [("İşyeri Kredisi", ISYERI_RE), ("Taşıt Kredisi", TASIT_RE),
                 ("İhtiyaç Kredisi", IHTIYAC_RE), ("Diğer", DIGER_RE)],
        "section": (re.compile(r"taksitli\s+ticari\s+kredi", re.IGNORECASE),
                    re.compile(r"kurumsal\s+kredi\s+kart", re.IGNORECASE)),
    },
}
DEFAULT_TABLES = ["tuketici", "taksitli"]


def field_order(spec):
    out = [("TP_Toplam", spec["label_tp"])]
    out += [(f"TP_S{i}", spec["subs"][i - 1][0]) for i in range(1, 5)]
    out += [("YP_Toplam", spec["label_yp"])]
    out += [(f"YP_S{i}", spec["subs"][i - 1][0]) for i in range(1, 5)]
    return out


def sums_ok(result):
    """True iff the four sub-items sum to the total, for both -TP and -YP."""
    for cur in ("TP", "YP"):
        parts = sum(result[f"{cur}_S{i}"] for i in range(1, 5))
        if result[f"{cur}_Toplam"] != parts:
            return False
    return True


MANUAL_OVERRIDES = {
    "tuketici": {
        # FİBABANKA A.Ş. — 30.09.2023, p.60
        (103, 2023, 9): {"TP_Toplam": 11119514, "TP_S1": 81896, "TP_S2": 498,
                         "TP_S3": 11037120, "TP_S4": 0,
                         "YP_Toplam": 0, "YP_S1": 0, "YP_S2": 0, "YP_S3": 0, "YP_S4": 0},
        # FİBABANKA A.Ş. — 31.03.2024, p.56
        (103, 2024, 3): {"TP_Toplam": 11806193, "TP_S1": 67580, "TP_S2": 484,
                         "TP_S3": 11738129, "TP_S4": 0,
                         "YP_Toplam": 0, "YP_S1": 0, "YP_S2": 0, "YP_S3": 0, "YP_S4": 0},
        # TÜRKİYE FİNANS KATILIM BANKASI A.Ş. — 31.03.2026, p.52
        (206, 2026, 3): {"TP_Toplam": 15012676, "TP_S1": 4589164, "TP_S2": 1597563,
                         "TP_S3": 8825949, "TP_S4": 0,
                         "YP_Toplam": 0, "YP_S1": 0, "YP_S2": 0, "YP_S3": 0, "YP_S4": 0},
        # TÜRKİYE EMLAK KATILIM BANKASI A.Ş. — 31.12.2025, p.93. This PDF draws
        # every figure one row BELOW its label, so a straight read yields a
        # blank -TP and shifted sub-items. The report's own grand total
        # (2.524.793 = 2.512.614 + 8.294 + 3.179 + 706) only reconciles under
        # the shifted reading, which these de-shifted figures use.
        (211, 2025, 12): {"TP_Toplam": 2512614, "TP_S1": 2457881, "TP_S2": 43373,
                          "TP_S3": 11360, "TP_S4": 0,
                          "YP_Toplam": 0, "YP_S1": 0, "YP_S2": 0, "YP_S3": 0, "YP_S4": 0},
    },
    "taksitli": {
        # TÜRKİYE FİNANS KATILIM BANKASI A.Ş. — 31.12.2024, p.100. Letter-spaced
        # Cari table, so OCR is the only route, and OCR drops the lone "1" in
        # İhtiyaç Kredileri (5,401 + 1 = 5,402). Read from the rendered page.
        (206, 2024, 12): {"TP_Toplam": 5402, "TP_S1": 0, "TP_S2": 5401,
                          "TP_S3": 1, "TP_S4": 0,
                          "YP_Toplam": 0, "YP_S1": 0, "YP_S2": 0, "YP_S3": 0, "YP_S4": 0},
    },
}


def assert_overrides_valid():
    for table, entries in MANUAL_OVERRIDES.items():
        for key, vals in entries.items():
            if set(vals) != set(FIELD_KEYS):
                raise AssertionError(f"Override {table} {key}: wrong field keys")
            if not sums_ok(vals):
                raise AssertionError(f"Override {table} {key}: sub-items do not sum to total")


WANTED_BANKS = [
    "AKBANK T.A.Ş.",
    "ALBARAKA TÜRK KATILIM BANKASI A.Ş.",
    "ALTERNATİFBANK A.Ş.",
    "ANADOLUBANK A.Ş.",
    "BURGAN BANK A.Ş.",
    "DENİZBANK A.Ş.",
    "FİBABANKA A.Ş.",
    "HAYAT FİNANS KATILIM BANKASI A.Ş.",
    "HSBC BANK A.Ş.",
    "ICBC TURKEY BANK A.Ş.",
    "ING BANK A.Ş.",
    "KUVEYT TÜRK KATILIM BANKASI A.Ş.",
    "QNB BANK A.Ş.",
    "T.C. ZİRAAT BANKASI A.Ş.",
    "TÜRK EKONOMİ BANKASI A.Ş.",
    "TÜRKİYE EMLAK KATILIM BANKASI A.Ş.",
    "TÜRKİYE FİNANS KATILIM BANKASI A.Ş.",
    "TÜRKİYE GARANTİ BANKASI A.Ş.",
    "TÜRKİYE HALK BANKASI A.Ş.",
    "TÜRKİYE VAKIFLAR BANKASI T.A.O.",
    "TÜRKİYE İŞ BANKASI A.Ş.",
    "VAKIF KATILIM BANKASI A.Ş.",
    "YAPI VE KREDİ BANKASI A.Ş.",
    "ZİRAAT KATILIM BANKASI A.Ş.",
    "ŞEKERBANK T.A.Ş.",
]

# Period markers — every report carries the current table AND a previous-period
# copy; only "Cari Dönem" may be used.
CARI_RE = re.compile(r"^CariD[öo]nem", re.IGNORECASE)
ONCEKI_RE = re.compile(r"^[ÖO]ncekiD[öo]nem", re.IGNORECASE)
# A period word alone is NOT enough: other notes in the same report have their
# own "Önceki Dönem ..." headers (e.g. "Önceki Dönem Ticari Tüketici", or the
# past-due aging table's "Önceki Dönem 31-60 Gün ... Toplam"). The header that
# governs a maturity-breakdown table always carries the vade columns, so
# require that too, and only look a few lines up (it sits 1-2 lines above).
VADELI_RE = re.compile(r"Vadeli", re.IGNORECASE)
PERIOD_LOOKBACK = 8

NUM_TOKEN_RE = re.compile(r"\(?-?\d[\d,.]*\)?|(?<!\S)-(?!\S)")


def make_session():
    session = requests.Session()
    session.headers.update(HEADERS)
    retry = Retry(total=5, backoff_factor=1.5, status_forcelist=[429, 500, 502, 503, 504])
    session.mount("https://", HTTPAdapter(max_retries=retry))
    session.mount("http://", HTTPAdapter(max_retries=retry))
    return session


def normalize_line(line):
    return line.replace("\xa0", " ").strip()


def parse_number(tok):
    tok = tok.strip()
    if tok in ("-", "–", "—", ""):
        return 0
    neg = tok.startswith("(") and tok.endswith(")")
    if neg:
        tok = tok[1:-1]
    if re.fullmatch(r"-?\d{1,3}(,\d{3})+(\.\d+)?", tok):
        tok = tok.replace(",", "")
    elif re.fullmatch(r"-?\d{1,3}(\.\d{3})+(,\d+)?", tok):
        tok = tok.replace(".", "").replace(",", ".")
    try:
        val = float(tok)
    except ValueError:
        return 0
    val = -val if neg else val
    return int(val) if val == int(val) else val


def line_last_number(line):
    toks = NUM_TOKEN_RE.findall(line)
    return parse_number(toks[-1]) if toks else None


def _period_of(keys, idx):
    """'cari' | 'onceki' | None — the period block the row at idx belongs to.

    Only a *governing* header counts (period word + this table's vade columns,
    a few lines above). None means no period header at all (e.g. Fibabanka),
    where the current table simply comes first.
    """
    for j in range(idx - 1, max(-1, idx - 1 - PERIOD_LOOKBACK), -1):
        key = keys[j]
        if not VADELI_RE.search(key):
            continue  # some other note's period header — not ours
        if ONCEKI_RE.match(key):
            return "onceki"
        if CARI_RE.match(key):
            return "cari"
    return None


def _pick_current_period_tp(keys, spec):
    """Index of the CURRENT period's -TP row; never a previous-period one."""
    cands = [i for i, k in enumerate(keys) if spec["hdr_tp"].match(k)]
    if not cands:
        return None, "TP header not found"
    marked = [(i, _period_of(keys, i)) for i in cands]
    for i, p in marked:
        if p == "cari":
            return i, "ok"
    for i, p in marked:
        if p is None:  # unmarked table (e.g. Fibabanka): current comes first
            return i, "ok"
    return None, "only Önceki Dönem table found (refusing previous-period data)"


def parse_lines(lines, spec):
    """Extract one table's 10 fields from 'label v1 v2 total' lines.

    Labels are matched on space-stripped keys; numbers are read from the
    original lines.
    """
    result = dict(ZERO_RESULT)
    n = len(lines)
    keys = [label_key(l) for l in lines]

    tp_idx, status = _pick_current_period_tp(keys, spec)
    if tp_idx is None:
        return result, status
    result["TP_Toplam"] = row_value(lines, tp_idx) or 0

    def find_after(start, pattern, window=8):
        for j in range(start + 1, min(start + 1 + window, n)):
            if pattern.match(keys[j]):
                return j
        return None

    idx = tp_idx
    for i, (_label, pat) in enumerate(spec["subs"], start=1):
        f = find_after(idx, pat)
        if f is not None:
            result[f"TP_S{i}"] = row_value(lines, f) or 0
            idx = f

    # Find -YP, skipping the "-Dövize Endeksli" block. Stop at "Önceki Dönem"
    # so a missing current-period YP can't fall through to the previous one.
    yp_idx = None
    for j in range(idx + 1, min(idx + 1 + 20, n)):
        if ONCEKI_RE.match(keys[j]):
            break
        if spec["hdr_yp"].match(keys[j]):
            yp_idx = j
            break
    if yp_idx is None:
        return result, "YP header not found"

    result["YP_Toplam"] = row_value(lines, yp_idx) or 0
    idx = yp_idx
    for i, (_label, pat) in enumerate(spec["subs"], start=1):
        f = find_after(idx, pat)
        if f is not None:
            result[f"YP_S{i}"] = row_value(lines, f) or 0
            idx = f

    return result, "ok"


def _pdf_to_lines(pdf_path):
    lines = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            lines.extend(normalize_line(l) for l in text.split("\n"))
    return lines


def _docx_to_lines(docx_path):
    doc = docx.Document(docx_path)
    lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_line(c.text) for c in row.cells]
            if cells and cells[0]:
                lines.append(" ".join(cells))
    return lines


def report_to_lines(report_path):
    if report_path.lower().endswith(".docx"):
        if docx is None:
            raise RuntimeError("python-docx not installed but report is .docx")
        return _docx_to_lines(report_path)
    return _pdf_to_lines(report_path)


# --------------------------------------------------------------------------
# OCR fallback (RapidOCR — pip-only, no system dependency)
# --------------------------------------------------------------------------
_OCR_ENGINE = None
_OCR_IMPORT_ERROR = None


def _get_ocr():
    global _OCR_ENGINE, _OCR_IMPORT_ERROR
    if _OCR_ENGINE is None and _OCR_IMPORT_ERROR is None:
        try:
            from rapidocr_onnxruntime import RapidOCR
            _OCR_ENGINE = RapidOCR()
        except Exception as e:  # noqa: BLE001
            _OCR_IMPORT_ERROR = e
    return _OCR_ENGINE


def _ocr_png_to_lines(png_path):
    """OCR a page image and rebuild table rows (cluster by y, order by x)."""
    engine = _get_ocr()
    if engine is None:
        raise RuntimeError(f"OCR engine unavailable: {_OCR_IMPORT_ERROR}")
    result, _ = engine(png_path)
    if not result:
        return []
    items = []
    for box, text, _score in result:
        ys = [p[1] for p in box]
        xs = [p[0] for p in box]
        items.append((sum(ys) / 4.0, min(xs), text))
    items.sort()
    rows, cur, last_y = [], [], None
    for y, x, txt in items:
        if last_y is None or abs(y - last_y) < 18:
            cur.append((x, txt))
        else:
            rows.append(sorted(cur))
            cur = [(x, txt)]
        last_y = y
    if cur:
        rows.append(sorted(cur))
    return [normalize_line(" ".join(t for _x, t in r)) for r in rows]


def _candidate_pages(pdf_path, spec):
    """Pages holding this table's section, plus the next (tables can spill).

    Requires both section-title phrases on the SAME page: that pinpoints the
    table page and avoids earlier prose pages that merely mention the topic.
    """
    sec1, sec2 = spec["section"]
    pages = set()
    try:
        with pdfplumber.open(pdf_path) as pdf:
            npages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                t = page.extract_text() or ""
                if sec1.search(t) and sec2.search(t):
                    pages.add(i)
                    if i + 1 < npages:
                        pages.add(i + 1)
    except Exception:
        return []
    return sorted(pages)[:4]


def extract_with_ocr(pdf_path, tag, ocr_dir, spec, dpi=300):
    """Render the section page(s) and OCR them. Accepts numbers only if the
    sum-check passes and the TP total is positive."""
    import fitz  # PyMuPDF

    candidates = _candidate_pages(pdf_path, spec)
    if not candidates:
        return dict(ZERO_RESULT), "ocr-no-section-page"

    os.makedirs(ocr_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    all_lines = []
    try:
        for pno in candidates:
            if pno >= len(doc):
                continue
            png = os.path.join(ocr_dir, f"{tag}_p{pno}.png")
            doc[pno].get_pixmap(dpi=dpi).save(png)
            all_lines.extend(_ocr_png_to_lines(png))
    finally:
        doc.close()

    result, status = parse_lines(all_lines, spec)
    if result["TP_Toplam"] <= 0:
        return result, f"ocr-parse-failed ({status}; page images saved)"
    if not sums_ok(result):
        return result, "ocr-sumcheck-failed (page images saved for manual read)"
    return result, "ok_ocr"


def extract_table(report_path, eft, year, month, ocr_dir, table, spec, lines=None):
    """Full chain for one bank/quarter/table: override -> text/docx -> OCR."""
    override = MANUAL_OVERRIDES.get(table, {}).get((eft, year, month))
    if override is not None:
        return dict(override), "manual_override"

    if lines is None:
        lines = report_to_lines(report_path)
    result, status = parse_lines(lines, spec)
    if status == "ok" and sums_ok(result):
        return result, "ok"

    # Either the block wasn't found, or its own arithmetic doesn't hold — the
    # latter means we mis-read something (e.g. TFKB's PDFs inject spaces inside
    # numbers: "2 3,640,299" reads back as 3,640,299). Rendering + OCR sidesteps
    # a broken text layer; we only trust the result if it sum-checks.
    tag = f"{table}_{eft}_{year}_{month:02d}"
    try:
        ocr_result, ocr_status = extract_with_ocr(report_path, tag, ocr_dir, spec)
    except Exception as e:  # noqa: BLE001
        if status == "ok":
            return result, f"needs_review: text sum-check failed, OCR unavailable ({e})"
        return dict(ZERO_RESULT), f"needs_review: text '{status}', OCR unavailable ({e})"
    if ocr_status == "ok_ocr":
        return ocr_result, "ok_ocr"
    if "no-section-page" in ocr_status and status != "ok":
        return dict(ZERO_RESULT), f"no_note ({status})"
    if status == "ok":
        # Found the block but its arithmetic doesn't hold and OCR couldn't
        # confirm a better reading -> emit 0 and flag, never unvalidated numbers.
        return dict(ZERO_RESULT), f"needs_review: sum-check failed, OCR could not confirm ({ocr_status})"
    return dict(ZERO_RESULT), f"needs_review: {ocr_status}"


def get_banks(session):
    r = session.post(f"{BASE}/Home/KurulusListesiGetir",
                     data={"kurulusTuruId": 1, "kurulusGrubuId": ""})
    r.raise_for_status()
    return {b["BankaAdi"]: b for b in r.json()}


RAPOR_URL_RE = re.compile(r'raporUrl=([^"&]+)')
KONSOLIDE_FILE_RE = re.compile(r"KONSOLIDE-(\d{4})-(\d{2})\.zip$", re.IGNORECASE)


def get_year_reports(session, eft_kodu, year):
    """{month: raporUrl} of KONSOLIDE reports for one bank/year."""
    params = {"KurulusTuru": 1, "EFTKodu": eft_kodu, "RaporTipi": "KONSOLIDE",
              "DonemYil": year, "DonemAy": 0}
    r = session.get(f"{BASE}/Home/SorguSonuc", params=params)
    r.raise_for_status()
    out = {}
    for raw_url in RAPOR_URL_RE.findall(r.text):
        url = raw_url.replace("&amp;", "&")
        m = KONSOLIDE_FILE_RE.search(requests.utils.unquote(url))
        if m:
            out[int(m.group(2))] = url
    return out


def _find_report_file(root_dir):
    candidates = []
    for root, _dirs, files in os.walk(root_dir):
        for fn in files:
            if fn.lower().endswith((".pdf", ".docx")):
                path = os.path.join(root, fn)
                candidates.append((os.path.getsize(path), path))
    return max(candidates)[1] if candidates else None


def _find_ext(root_dir, ext):
    for root, _dirs, files in os.walk(root_dir):
        for fn in files:
            if fn.lower().endswith(ext):
                return os.path.join(root, fn)
    return None


def download_report(session, rapor_url, cache_dir, cache_key):
    """Download+extract the report (cached), return path to the PDF/DOCX.
    Handles raw-PDF-as-zip, zip-within-zip, and multi-file archives."""
    os.makedirs(cache_dir, exist_ok=True)
    rep_dir = os.path.join(cache_dir, cache_key)
    marker = os.path.join(rep_dir, ".report_path")
    if os.path.exists(marker):
        with open(marker) as f:
            p = f.read().strip()
        if p and os.path.exists(p):
            return p

    decoded_url = urllib.parse.unquote(rapor_url)
    r = session.get(f"{BASE}/Home/DosyaIndir", params={"raporUrl": decoded_url}, timeout=120)
    r.raise_for_status()
    os.makedirs(rep_dir, exist_ok=True)

    if r.content[:4] == b"%PDF":
        report_path = os.path.join(rep_dir, "report.pdf")
        with open(report_path, "wb") as f:
            f.write(r.content)
    else:
        with zipfile.ZipFile(io.BytesIO(r.content)) as z:
            z.extractall(rep_dir)
        for _ in range(3):
            report_path = _find_report_file(rep_dir)
            if report_path:
                break
            nested = _find_ext(rep_dir, ".zip")
            if not nested:
                break
            with zipfile.ZipFile(nested) as z:
                z.extractall(rep_dir)
        report_path = _find_report_file(rep_dir)

    if report_path is None:
        raise RuntimeError(f"No PDF/DOCX report found inside archive for {rapor_url}")
    with open(marker, "w") as f:
        f.write(report_path)
    return report_path


# --------------------------------------------------------------------------
# Latest-quarter detection
# --------------------------------------------------------------------------
def q_label(year, month):
    return f"{year}Q{month // 3}"


def parse_quarter_arg(s):
    m = re.fullmatch(r"(\d{4})Q([1-4])", s.strip(), re.IGNORECASE)
    if not m:
        raise SystemExit(f"--quarter must look like 2026Q2 (got {s!r})")
    return int(m.group(1)), int(m.group(2)) * 3


def gather_reports(session, banks, years, delay):
    """{eft: {(year, month): url}} for the given years."""
    out = {}
    for name in WANTED_BANKS:
        eft = banks[name]["EFTKodu"]
        d = {}
        for yr in years:
            try:
                for m, u in get_year_reports(session, eft, yr).items():
                    d[(yr, m)] = u
            except Exception as e:  # noqa: BLE001
                print(f"  WARN: report index fetch failed for {name} {yr}: {e}")
            time.sleep(delay)
        out[eft] = d
    return out


# --------------------------------------------------------------------------
# Workbook — one sheet per table, Template.xlsx layout
# --------------------------------------------------------------------------
LABEL_COL = 6   # F
FIRST_Q_COL = 7  # G
HEADER_FONT = openpyxl.styles.Font(name="Helvetica", size=12, bold=True)
QUARTER_FONT = openpyxl.styles.Font(name="Helvetica", size=12, bold=False)
FIELD_FONT = openpyxl.styles.Font(name="Helvetica", size=12, bold=False)
YELLOW_FILL = openpyxl.styles.PatternFill(fill_type="solid", fgColor="FFFFFF00")
VALUE_ALIGN = openpyxl.styles.Alignment(horizontal="left")


def _write_sheet(ws, spec, quarters, bank_results, bank_order):
    row = 1
    first_bank = True
    for bank_name in bank_order:
        name_cell = ws.cell(row=row, column=LABEL_COL, value=bank_name)
        name_cell.font = HEADER_FONT
        name_cell.fill = YELLOW_FILL
        if first_bank:
            for j, (y, q3, _lbl) in enumerate(quarters):
                qcell = ws.cell(row=row, column=FIRST_Q_COL + j, value=f"{y} Q{q3 // 3}")
                qcell.font = QUARTER_FONT
                qcell.fill = YELLOW_FILL
            first_bank = False
        row += 1
        for key, label in field_order(spec):
            ws.cell(row=row, column=LABEL_COL, value=label).font = FIELD_FONT
            for j, (_y, _m, qlabel) in enumerate(quarters):
                val = bank_results.get(bank_name, {}).get(qlabel, ZERO_RESULT)[key]
                vcell = ws.cell(row=row, column=FIRST_Q_COL + j, value=val)
                vcell.font = FIELD_FONT
                vcell.alignment = VALUE_ALIGN
                vcell.number_format = "#,##0"
            row += 1
    ws.column_dimensions[openpyxl.utils.get_column_letter(LABEL_COL)].width = 35.5
    for j in range(len(quarters)):
        ws.column_dimensions[openpyxl.utils.get_column_letter(FIRST_Q_COL + j)].width = 17
    ws.freeze_panes = ws.cell(row=2, column=FIRST_Q_COL).coordinate


def build_workbook(tables, quarters, results_by_table, bank_order):
    """One workbook, one sheet per table (their sub-items differ)."""
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for table in tables:
        spec = TABLE_SPECS[table]
        ws = wb.create_sheet(title=spec["sheet"])
        _write_sheet(ws, spec, quarters, results_by_table[table], bank_order)
    return wb


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--quarter", default=None,
                    help="Pin a quarter, e.g. 2026Q2. Default: auto-detect newest.")
    ap.add_argument("--tables", default=",".join(DEFAULT_TABLES),
                    help="Comma list: tuketici,taksitli (default: both -> two sheets)")
    ap.add_argument("--out-dir", default="output")
    ap.add_argument("--cache-dir", default="bddk_cache")
    ap.add_argument("--ocr-dir", default="_ocr_pages")
    ap.add_argument("--delay", type=float, default=0.4)
    args = ap.parse_args()

    tables = [t.strip() for t in args.tables.split(",") if t.strip()]
    for t in tables:
        if t not in TABLE_SPECS:
            raise SystemExit(f"Unknown table {t!r}; choose from {list(TABLE_SPECS)}")
    assert_overrides_valid()

    session = make_session()
    print("Fetching bank list...")
    banks = get_banks(session)
    missing = [n for n in WANTED_BANKS if n not in banks]
    if missing:
        raise SystemExit("WANTED_BANKS not found in bank list: " + "; ".join(missing))

    if args.quarter:
        year, month = parse_quarter_arg(args.quarter)
        print(f"Using pinned quarter {q_label(year, month)}.")
        reports = gather_reports(session, banks, [year], args.delay)
    else:
        cur = datetime.now().year
        print(f"Detecting latest available quarter (scanning {cur-1}-{cur})...")
        reports = gather_reports(session, banks, [cur, cur - 1], args.delay)
        all_ym = set().union(*[set(d) for d in reports.values()]) if reports else set()
        if not all_ym:
            raise SystemExit("No consolidated reports found for any of the 25 banks.")
        year, month = max(all_ym)

    target = (year, month)
    have = sum(1 for eft_map in reports.values() if target in eft_map)
    qlbl = q_label(year, month)
    print(f"\nLatest quarter: {qlbl}  —  {have}/{len(WANTED_BANKS)} banks have a report.")
    if have < len(WANTED_BANKS) / 2:
        print("  WARNING: fewer than half the banks have filed this quarter yet — "
              "filing may still be in progress; consider re-running later or pass "
              "--quarter for an earlier one.")

    quarters = [(year, month, qlbl)]
    os.makedirs(args.out_dir, exist_ok=True)
    log_path = os.path.join(args.out_dir, f"run_{qlbl}.jsonl")
    out_path = os.path.join(args.out_dir, f"BDDK_25Banks_{qlbl}.xlsx")

    results = {t: {} for t in tables}
    summary = {t: {"ok": 0, "ok_ocr": 0, "manual_override": 0, "no_note": 0}
               for t in tables}
    no_report = 0
    needs_review, errors = [], []

    with open(log_path, "w", encoding="utf-8") as logf:
        for i, name in enumerate(WANTED_BANKS, 1):
            eft = banks[name]["EFTKodu"]
            print(f"[{i}/{len(WANTED_BANKS)}] {name} (EFT {eft})")
            rapor_url = reports[eft].get(target)
            if not rapor_url:
                for t in tables:
                    results[t][name] = {qlbl: dict(ZERO_RESULT)}
                no_report += 1
                logf.write(json.dumps({"bank": name, "eft": eft, "quarter": qlbl,
                                       "table": "*", "status": "no_consolidated_report"},
                                      ensure_ascii=False) + "\n")
                continue
            try:
                cache_key = f"{eft}_{year}_{month:02d}"
                report_path = download_report(session, rapor_url, args.cache_dir, cache_key)
                # Read the report once, parse every requested table from it.
                lines = report_to_lines(report_path)
                for t in tables:
                    res, status = extract_table(report_path, eft, year, month,
                                                args.ocr_dir, t, TABLE_SPECS[t], lines=lines)
                    results[t][name] = {qlbl: res}
                    logf.write(json.dumps({"bank": name, "eft": eft, "quarter": qlbl,
                                           "table": t, "status": status},
                                          ensure_ascii=False) + "\n")
                    if status in summary[t]:
                        summary[t][status] += 1
                    elif status.startswith("no_note"):
                        summary[t]["no_note"] += 1
                    elif status.startswith("needs_review"):
                        needs_review.append((name, t, status))
                        print(f"  NEEDS REVIEW [{t}]: {status}")
                    if status == "ok_ocr":
                        print(f"  OCR [{t}]: read from image, sum-check passed")
            except Exception as e:  # noqa: BLE001
                for t in tables:
                    results[t][name] = {qlbl: dict(ZERO_RESULT)}
                errors.append((name, str(e)))
                logf.write(json.dumps({"bank": name, "eft": eft, "quarter": qlbl,
                                       "table": "*", "status": f"error: {e}"},
                                      ensure_ascii=False) + "\n")
                print(f"  ERROR {qlbl}: {e} -> 0")
            time.sleep(args.delay)

    build_workbook(tables, quarters, results, WANTED_BANKS).save(out_path)

    print("\n" + "=" * 62)
    print(f"Quarter {qlbl}: wrote {out_path}")
    print(f"  sheets: {', '.join(TABLE_SPECS[t]['sheet'] for t in tables)}")
    for t in tables:
        s = summary[t]
        print(f"  [{TABLE_SPECS[t]['sheet']}] text-parsed: {s['ok']}  OCR: {s['ok_ocr']}  "
              f"override: {s['manual_override']}  no-note(0): {s['no_note']}")
    print(f"  banks with no report for {qlbl} (all 0): {no_report}")
    if needs_review:
        print(f"  NEEDS MANUAL REVIEW ({len(needs_review)}) — images in {args.ocr_dir}/:")
        for name, t, st in needs_review:
            print(f"    - {name} [{t}]: {st}")
    if errors:
        print(f"  ERRORS ({len(errors)}) — left as 0 (usually a file missing on BDDK):")
        for name, e in errors:
            print(f"    - {name}: {e}")
    print(f"  log: {log_path}")


if __name__ == "__main__":
    main()
